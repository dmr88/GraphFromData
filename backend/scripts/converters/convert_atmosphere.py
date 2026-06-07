#!/usr/bin/env python3
"""
Чистит таблицы из docx-файлов «атмосфера ВКО»:
- ИЗА5 по годам (2014-2024) для промышленных городов
- Общий объём выбросов от стационарных источников
- Газообразные / твёрдые выбросы
- Объёмы PM10, PM2.5, SO2, H2S, CO, NOx, ЛОС по городам (РМ и элементы)

Результат: public/Атмосфера_ВКО_очищенный.xlsx
"""
import os
import re
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment

BASE = '/sessions/inspiring-peaceful-bell/mnt/GraphFromData/public/атмосфера ВКО'
DST  = '/sessions/inspiring-peaceful-bell/mnt/GraphFromData/public/Атмосфера_ВКО_очищенный.xlsx'


def parse_number(s):
    if s is None:
        return None
    s = str(s).strip().replace(',', '.').replace(' ', '').replace(' ', '')
    if not s or s in ('-', '—'):
        return None
    try:
        return float(s)
    except ValueError:
        # Может быть "1,5 (2)" или "0,8±0,1" — берём первое число
        m = re.match(r'^(-?\d*\.?\d+)', s)
        if m:
            try:
                return float(m.group(1))
            except ValueError:
                return None
        return None


# Унификация написания населённых пунктов
CITY_ALIASES = {
    'оскемен': 'Өскемен',
    'усть-каменогорск': 'Өскемен',
    'г. усть-каменогорск': 'Өскемен',
    'г.усть-каменогорск': 'Өскемен',
    'өскемен': 'Өскемен',
    'риддер': 'Риддер',
    'г.риддер': 'Риддер',
    'семей': 'Семей',
    'г.семей': 'Семей',
    'глубокое': 'Глубокое',
    'глубоковское': 'Глубокое',
    'п.глубокое': 'Глубокое',
    'алтай': 'Алтай',
    'г.алтай': 'Алтай',
    'алтай (зырян)': 'Алтай',
    'зырян': 'Алтай',
    'шемонаиха': 'Шемонаиха',
    'г.шемонаиха': 'Шемонаиха',
    'курчатов': 'Курчатов',
    'г.курчатов': 'Курчатов',
    'абай': 'Абай',
    'аягөз': 'Аягөз',
    'аягоз': 'Аягөз',
    'бесқарағай': 'Бесқарағай',
    'бескарагай': 'Бесқарағай',
    'бородулиха': 'Бородулиха',
    'жарма': 'Жарма',
    'зайсан': 'Зайсан',
    'катон-қарағай': 'Катон-Қарағай',
    'катон-карагай': 'Катон-Қарағай',
    'көкпекті': 'Көкпекті',
    'кокпекты': 'Көкпекті',
    'кокпектинский': 'Көкпекті',
    'күршім': 'Күршім',
    'курчум': 'Күршім',
    'курчим': 'Күршім',
    'курчумский': 'Күршім',
    'тарбағатай': 'Тарбағатай',
    'тарбагатай': 'Тарбағатай',
    'ұлан': 'Ұлан',
    'улан': 'Ұлан',
    'үржар': 'Үржар',
    'урджар': 'Үржар',
}


def norm_city(name):
    s = ' '.join(str(name).split()).strip()
    return CITY_ALIASES.get(s.lower(), s)


def style_header(ws):
    f = Font(bold=True, color='FFFFFF')
    fill = PatternFill('solid', start_color='2563EB')
    for cell in ws[1]:
        cell.font = f
        cell.fill = fill
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    ws.row_dimensions[1].height = 32


def autosize(ws):
    for col in ws.columns:
        m = 0
        letter = col[0].column_letter
        for c in col:
            v = str(c.value) if c.value is not None else ''
            m = max(m, min(len(v), 40))
        ws.column_dimensions[letter].width = max(10, m + 2)


def extract_year_city_table(table):
    """Извлекает таблицу формата: первая строка — заголовки (город, год1, год2, ...),
    последующие — город + значения. Возвращает (years, rows_dict, extra_col_name).
    """
    rows = table.rows
    header = [c.text.strip() for c in rows[0].cells]
    years = []
    year_indices = []
    extra_col_idx = None
    extra_col_name = None
    for i, h in enumerate(header[1:], start=1):
        m = re.match(r'^\s*(\d{4})\s*$', h)
        if m:
            years.append(int(m.group(1)))
            year_indices.append(i)
        elif h:
            extra_col_idx = i
            extra_col_name = h

    data = {}
    extras = {}
    for r in rows[1:]:
        cells = [c.text.strip() for c in r.cells]
        if not cells[0]:
            continue
        city = norm_city(cells[0])
        if not city or city.lower() in ('населенный пункт', 'населённый пункт'):
            continue
        values = []
        for idx in year_indices:
            v = parse_number(cells[idx] if idx < len(cells) else None)
            values.append(v)
        data[city] = values
        if extra_col_idx is not None and extra_col_idx < len(cells):
            extras[city] = cells[extra_col_idx]
    return years, data, extras, extra_col_name


def add_year_city_sheet(wb, name, years, data, extras=None, extra_col=None):
    ws = wb.create_sheet(name[:31])
    header = ['Город'] + [str(y) for y in years]
    if extra_col:
        header.append(extra_col)
    ws.append(header)
    for city, vals in data.items():
        row = [city] + vals
        if extra_col:
            row.append(extras.get(city, ''))
        ws.append(row)
    style_header(ws)
    autosize(ws)


def main():
    out = Workbook()
    out.remove(out.active)
    long_records = []  # (indicator, year, city, value)

    # ===== 1. ИЗА5 (наиболее полная — из 2024/ИЗА 2024.docx) =====
    doc = Document(os.path.join(BASE, '2024', 'ИЗА 2024.docx'))
    table = doc.tables[0]
    # Первая строка — повторение заголовка ('ИЗА5'), вторая — годы
    # Используем вторую строку как заголовок
    header_row = [c.text.strip() for c in table.rows[1].cells]
    years = []
    year_indices = []
    extra_idx = None
    extra_name = None
    for i, h in enumerate(header_row[1:], start=1):
        m = re.match(r'^\s*(\d{4})\s*$', h)
        if m:
            years.append(int(m.group(1)))
            year_indices.append(i)
        elif h and 'отрасль' in h.lower():
            extra_idx = i
            extra_name = 'Отрасль промышленности'

    data, extras = {}, {}
    for r in table.rows[2:]:
        cells = [c.text.strip() for c in r.cells]
        city = norm_city(cells[0])
        if not city:
            continue
        vals = [parse_number(cells[i] if i < len(cells) else None) for i in year_indices]
        data[city] = vals
        if extra_idx is not None and extra_idx < len(cells):
            extras[city] = cells[extra_idx]
    add_year_city_sheet(out, 'ИЗА5', years, data, extras, extra_name)
    for city, vals in data.items():
        for y, v in zip(years, vals):
            if v is not None:
                long_records.append(('ИЗА5', y, city, v))

    # ===== 2. Общий объём выбросов от стац. источников =====
    # data из «данные от стац ист.docx» Table 0 (2014-2023) + из 2024/общий объем 2024.docx (2024)
    doc1 = Document(os.path.join(BASE, 'данные от стац ист.docx'))
    years_main, data_main, _, _ = extract_year_city_table(doc1.tables[0])
    doc2 = Document(os.path.join(BASE, '2024', 'общий объем 2024.docx'))
    years_2024, data_2024, _, _ = extract_year_city_table(doc2.tables[0])

    all_years = years_main + [y for y in years_2024 if y not in years_main]
    cities = list(dict.fromkeys(list(data_main.keys()) + list(data_2024.keys())))
    merged = {}
    for c in cities:
        row = []
        for y in all_years:
            v = None
            if y in years_main and c in data_main:
                v = data_main[c][years_main.index(y)]
            if v is None and y in years_2024 and c in data_2024:
                v = data_2024[c][years_2024.index(y)]
            row.append(v)
        merged[c] = row
    add_year_city_sheet(out, 'Общий объём выбросов', all_years, merged)
    for c, vals in merged.items():
        for y, v in zip(all_years, vals):
            if v is not None:
                long_records.append(('Общий объём выбросов, тыс. т', y, c, v))

    # ===== 3. Газообразные выбросы =====
    doc1 = Document(os.path.join(BASE, 'газобраз.docx'))
    # Первая таблица 12 строк (промышленные), вторая 21 (все) — берём более полную (21)
    y1, d1, _, _ = extract_year_city_table(doc1.tables[1])
    doc2 = Document(os.path.join(BASE, '2024', 'газообразные 2024.docx'))
    y2, d2, _, _ = extract_year_city_table(doc2.tables[0])
    all_years = y1 + [y for y in y2 if y not in y1]
    cities = list(dict.fromkeys(list(d1.keys()) + list(d2.keys())))
    merged = {}
    for c in cities:
        row = []
        for y in all_years:
            v = None
            if y in y1 and c in d1:
                v = d1[c][y1.index(y)]
            if v is None and y in y2 and c in d2:
                v = d2[c][y2.index(y)]
            row.append(v)
        merged[c] = row
    add_year_city_sheet(out, 'Газообразные выбросы', all_years, merged)
    for c, vals in merged.items():
        for y, v in zip(all_years, vals):
            if v is not None:
                long_records.append(('Газообразные выбросы, тыс. т', y, c, v))

    # ===== 4. Твёрдые выбросы =====
    doc = Document(os.path.join(BASE, '2024', 'твердые 2024.docx'))
    y2024, d2024, _, _ = extract_year_city_table(doc.tables[0])
    add_year_city_sheet(out, 'Твёрдые выбросы (2024)', y2024, d2024)
    for c, vals in d2024.items():
        for y, v in zip(y2024, vals):
            if v is not None:
                long_records.append(('Твёрдые выбросы, тыс. т', y, c, v))

    # ===== 5. PM10, PM2.5, SO2, H2S, CO, NOx, ЛОС по городам (РМ и элементы.docx) =====
    doc = Document(os.path.join(BASE, '2024', 'РМ и элементы.docx'))
    # Каждая таблица — отдельный город. В первой ячейке заголовка — название города.
    # Колонки: 'Город' / 2014 / 2015 / ... / 2025
    # Строки: каждое вещество.
    pollutants = ['PM10', 'PM2.5', 'SO2', 'H2S', 'CO', 'NOx', 'ЛОС']
    pollutant_data = {p: {} for p in pollutants}  # вещество → {город: {год: значение}}
    pollutant_years = set()
    for table in doc.tables:
        header = [c.text.strip() for c in table.rows[0].cells]
        city = norm_city(header[0])
        if not city:
            continue
        years_p = []
        year_idx_p = []
        for i, h in enumerate(header[1:], start=1):
            m = re.match(r'^\s*(\d{4})\s*$', h)
            if m:
                years_p.append(int(m.group(1)))
                year_idx_p.append(i)
                pollutant_years.add(int(m.group(1)))
        for r in table.rows[1:]:
            cells = [c.text.strip() for c in r.cells]
            name = cells[0].lower()
            target = None
            if 'pm10' in name: target = 'PM10'
            elif 'pm2' in name: target = 'PM2.5'
            elif 'so2' in name or 'сернистый' in name: target = 'SO2'
            elif 'h2s' in name or 'сероводород' in name: target = 'H2S'
            elif 'co)' in name or 'углерода' in name: target = 'CO'
            elif 'азот' in name or 'nox' in name: target = 'NOx'
            elif 'летучие' in name or 'лос' in name: target = 'ЛОС'
            if not target:
                continue
            for idx, y in zip(year_idx_p, years_p):
                v = parse_number(cells[idx] if idx < len(cells) else None)
                if v is None:
                    continue
                pollutant_data[target].setdefault(city, {})[y] = v
                long_records.append((f'Выброс {target}, тыс. т', y, city, v))

    sorted_years = sorted(pollutant_years)
    for poll in pollutants:
        cities_p = list(pollutant_data[poll].keys())
        if not cities_p:
            continue
        data_p = {c: [pollutant_data[poll][c].get(y) for y in sorted_years] for c in cities_p}
        add_year_city_sheet(out, f'Выброс {poll}', sorted_years, data_p)

    # ===== Лист «Полные данные» =====
    ws_full = out.create_sheet('Полные данные')
    ws_full.append(['Год', 'Показатель', 'Город', 'Значение'])
    for indicator, year, city, value in long_records:
        ws_full.append([str(year), indicator, city, value])
    style_header(ws_full)
    autosize(ws_full)
    ws_full.freeze_panes = 'A2'

    out.save(DST)
    print(f"[OK] Сохранено: {DST}")
    print(f"     Листов: {len(out.sheetnames)}, длинных записей: {len(long_records)}")
    for n in out.sheetnames:
        ws_n = out[n]
        print(f"  - {n}: {ws_n.max_row-1} строк × {ws_n.max_column} кол")


if __name__ == '__main__':
    main()
